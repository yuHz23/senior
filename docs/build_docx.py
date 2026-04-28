import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def add_centered_para(doc, text, bold=False, size=12, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def build_cover_pages(doc):
    # Page 1
    add_centered_para(doc, "VIETNAM NATIONAL UNIVERSITY – HO CHI MINH CITY", False, 12)
    add_centered_para(doc, "INTERNATIONAL UNIVERSITY", False, 12)
    add_centered_para(doc, "SCHOOL OF ELECTRICAL ENGINEERING", False, 12, 100)
    
    add_centered_para(doc, "NGHIÊN CỨU ỨNG DỤNG LSTM AUTOENCODER TRONG NÉN DỮ LIỆU SEMANTIC VÀ PHÁT HIỆN ANOMALY TRÊN IOT GATEWAY", True, 16)
    add_centered_para(doc, "(RESEARCH ON THE APPLICATION OF LSTM AUTOENCODER IN SEMANTIC DATA COMPRESSION AND ANOMALY DETECTION ON IOT GATEWAY)", True, 14, 100)
    
    add_centered_para(doc, "BY", False, 12)
    add_centered_para(doc, "NGO HOANG HUY", True, 12)
    add_centered_para(doc, "EEEEIU21019", True, 12, 80)
    
    add_centered_para(doc, "A SENIOR PROJECT SUBMITTED TO THE SCHOOL OF ELECTRICAL ENGINEERING", False, 12)
    add_centered_para(doc, "IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR", False, 12)
    add_centered_para(doc, "THE DEGREE OF ENGINEER IN ELECTRONICS – TELECOMMUNICATIONS ENGINEERING", False, 12, 100)
    
    add_centered_para(doc, "HO CHI MINH CITY, VIETNAM", False, 12)
    add_centered_para(doc, "April 2026", False, 12)
    
    doc.add_page_break()
    
    # Page 2
    add_centered_para(doc, "i RESEARCH ON THE APPLICATION OF LSTM AUTOENCODER IN SEMANTIC DATA COMPRESSION AND ANOMALY DETECTION ON IOT GATEWAY", True, 14, 40)
    add_centered_para(doc, "BY", False, 12, 20)
    add_centered_para(doc, "NGO HOANG HUY\nEEEEIU21019", True, 12, 40)
    
    p = doc.add_paragraph("Under the guidance and approval of the committee, and approved by its members, this senior project has been accepted in partial fulfillment of the requirements for the degree.")
    p.paragraph_format.space_after = Pt(30)
    
    p = doc.add_paragraph("Approved by:\n")
    
    p = doc.add_paragraph("________________________________\nChairperson")
    p.paragraph_format.space_before = Pt(30)
    
    p = doc.add_paragraph("________________________________\nCommittee member")
    p.paragraph_format.space_before = Pt(30)
    
    p = doc.add_paragraph("________________________________      ________________________________\nCommittee member                                        Committee member")
    p.paragraph_format.space_before = Pt(30)
    
    p = doc.add_paragraph("________________________________\nCommittee member")
    p.paragraph_format.space_before = Pt(30)
    
    doc.add_page_break()

def add_md_to_docx(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Simple markdown parser
    # Remove frontmatter
    content = re.sub(r'^---.*?---\n', '', content, flags=re.DOTALL)
    
    blocks = content.split('\n\n')
    skip_next = False
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        
        # Skip the duplicate title blocks since we already added cover pages
        if "VIETNAM NATIONAL UNIVERSITY" in block or "INTERNATIONAL UNIVERSITY" in block or "SCHOOL OF ELECTRICAL ENGINEERING" in block:
            continue
        if "Nghiên cứu ứng dụng" in block or "(Research on the Application" in block or "NGO HOANG HUY" in block or "HO CHI MINH CITY" in block or "A SENIOR PROJECT" in block:
            continue
        if block == "---" or block == '<div style="page-break-after: always;"></div>':
            doc.add_page_break()
            continue
            
        if block.startswith('# '):
            doc.add_heading(block[2:], level=1)
        elif block.startswith('## '):
            doc.add_heading(block[3:], level=2)
        elif block.startswith('### '):
            doc.add_heading(block[4:], level=3)
        elif block.startswith('!['):
            # Extract image path
            match = re.search(r'\!\[.*?\]\((.*?)\)', block)
            if match:
                img_path = match.group(1)
                img_path = img_path.replace('./', '')
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6.0))
        else:
            p = doc.add_paragraph()
            # Handle basic bold parsing
            parts = re.split(r'(\*\*.*?\*\*)', block)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    # also handle italics * *
                    it_parts = re.split(r'(\*.*?\*)', part)
                    for it_part in it_parts:
                        if it_part.startswith('*') and it_part.endswith('*'):
                            p.add_run(it_part[1:-1]).italic = True
                        else:
                            p.add_run(it_part)

def main():
    base_docx = r"D:\Sen_Claude\Report_NgoHoangHuy_EEEEIU21019.docx"
    doc = docx.Document(base_docx)
    
    # Clear existing content to retain just styles/margins/headers
    doc._body.clear_content()
    
    # Set default font
    if 'Normal' in doc.styles:
        style = doc.styles['Normal']
        if style.font:
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
    
    build_cover_pages(doc)
    
    md_path = "Report_NgoHoangHuy_EEEEIU21019_Final.md"
    add_md_to_docx(doc, md_path)
    
    output_docx = r"D:\Sen_Claude\Report_NgoHoangHuy_EEEEIU21019_Final.docx"
    doc.save(output_docx)
    print(f"Successfully generated {output_docx}")

if __name__ == "__main__":
    main()
